import sys
import os
sys.path.insert(0, os.path.abspath(os.path.dirname(__file__)))

from docx import Document
from app.services.report_analyzer import groq_service
import asyncio

# Create a simple test DOCX
test_docx_path = "test_dental_report.docx"
doc = Document()
doc.add_heading("Dental Report", level=1)
doc.add_paragraph("Patient Name: John Doe")
doc.add_paragraph("Age: 30")
doc.add_paragraph("Date: 2025-01-01")
doc.add_paragraph("Clinic: Smilo Dental")
doc.add_paragraph("Findings: Mild gum inflammation, cavity in tooth 14")
doc.add_paragraph("Recommendations: Scaling and filling")
doc.save(test_docx_path)
print(f"Created test DOCX: {test_docx_path}")

# Read it
with open(test_docx_path, "rb") as f:
    file_content = f.read()

async def test_it():
    print("\n=== Testing Report Analyzer ===")
    result = await groq_service.analyze_dental_report(
        file_content, "test_dental_report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    print("\n=== RESULT ===")
    print(result.model_dump_json(indent=2))

asyncio.run(test_it())
